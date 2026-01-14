"""
Resume Parser - PDF and DOCX text extraction.

Implements FE-03: Extract text from resume files for AI context.
Uses PyMuPDF (fitz) for PDF and python-docx for DOCX.
"""

import re
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document


class ResumeParser:
    """
    Parser for extracting text from resume files.
    
    Supports PDF and DOCX formats.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    @classmethod
    def is_supported(cls, file_path: Path) -> bool:
        """Check if file type is supported."""
        return file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def extract_text(cls, file_path: Path) -> str:
        """
        Extract text from a resume file.
        
        Args:
            file_path: Path to the PDF or DOCX file.
            
        Returns:
            Extracted text content.
            
        Raises:
            ValueError: If file type is not supported.
            FileNotFoundError: If file doesn't exist.
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return cls._extract_pdf(file_path)
        elif suffix in {".docx", ".doc"}:
            return cls._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        text_parts = []
        
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                text = page.get_text()
                if text:
                    text_parts.append(text)
        
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        doc = Document(str(file_path))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)

    @classmethod
    def extract_contact_info(cls, text: str) -> dict[str, Optional[str]]:
        """
        Extract basic contact information from resume text.
        
        Args:
            text: Resume text content.
            
        Returns:
            Dictionary with extracted email, phone, and name.
        """
        info: dict[str, Optional[str]] = {
            "email": None,
            "phone": None,
            "name": None,
        }
        
        # Email pattern
        email_match = re.search(
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            text
        )
        if email_match:
            info["email"] = email_match.group()
        
        # Phone pattern (various formats)
        phone_match = re.search(
            r'(?:\+?55\s?)?(?:\(?0?\d{2}\)?\s?)?(?:9\s?)?\d{4}[-.\s]?\d{4}',
            text
        )
        if phone_match:
            info["phone"] = phone_match.group()
        
        # Name: usually first non-empty line
        lines = text.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Name is typically just letters and spaces
            if line and re.match(r'^[A-Za-zÀ-ÿ\s]+$', line) and len(line) > 3:
                info["name"] = line
                break
        
        return info

    @classmethod
    def extract_skills(cls, text: str) -> list[str]:
        """
        Extract likely skills from resume text.
        
        Args:
            text: Resume text content.
            
        Returns:
            List of identified skills.
        """
        # Common tech skills to look for
        common_skills = [
            "Python", "Java", "JavaScript", "TypeScript", "C#", "C++",
            "React", "Angular", "Vue", "Node.js", "Django", "Flask",
            "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis",
            "AWS", "Azure", "GCP", "Docker", "Kubernetes",
            "Git", "Linux", "CI/CD", "REST", "GraphQL",
            "Scrum", "Agile", "TDD", "Clean Code",
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in common_skills:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        
        return found_skills
